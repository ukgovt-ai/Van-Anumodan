"""
Van Anumodan — desktop extraction engine.
Reads every proposal document format (including SCANNED PDFs and images via OCR,
English + Hindi/Devanagari) and returns structured figures for scrutiny.
All processing is local; nothing leaves the machine.
"""
import io, re, os, sys, zipfile, json, math

def _configure_tesseract():
    """Point pytesseract at a bundled or installed Tesseract, and its tessdata."""
    import pytesseract
    cand = os.environ.get("VA_TESSERACT")
    # bundled inside a PyInstaller exe: <bundle>/tesseract/tesseract.exe
    if hasattr(sys, "_MEIPASS"):
        b = os.path.join(sys._MEIPASS, "tesseract", "tesseract.exe")
        if os.path.isfile(b):
            cand = b
    # common Windows install location
    if not cand:
        for p in (r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                  r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"):
            if os.path.isfile(p):
                cand = p; break
    if cand and os.path.isfile(cand):
        pytesseract.pytesseract.tesseract_cmd = cand
        td = os.path.join(os.path.dirname(cand), "tessdata")
        if os.path.isdir(td):
            os.environ.setdefault("TESSDATA_PREFIX", os.path.dirname(cand))


# Optional heavy deps are imported lazily so the module can be inspected without them.
def _lazy():
    global fitz, pytesseract, Image, openpyxl, docx
    import fitz  # PyMuPDF
    import pytesseract
    from PIL import Image
    import openpyxl
    import docx
    _configure_tesseract()
    return fitz, pytesseract, Image, openpyxl, docx

OCR_LANG = os.environ.get("VA_OCR_LANG", "eng+hin")
OCR_DPI = int(os.environ.get("VA_OCR_DPI", "300"))
TEXT_LAYER_MIN = 40  # chars; below this a PDF is treated as scanned

# ---------- figure extraction (English + Hindi) ----------
def _num(s):
    try:
        return float(str(s).replace(",", "").strip())
    except Exception:
        return None

def _find_area(t, kw):
    m = re.search(kw + r"[^0-9]{0,40}(\d{1,3}(?:,\d{3})*(?:\.\d+)?)\s*(?:ha|hectare|hect|\u0939\u0947\u0915\u094d\u091f\u0947\u092f\u0930)", t, re.I)
    return _num(m.group(1)) if m else None

def _find_count(t, kw):
    m = re.search(kw + r"[^0-9]{0,30}(\d{1,3}(?:,\d{3})*)", t, re.I)
    return _num(m.group(1)) if m else None

def extract_figures(text):
    t = re.sub(r"\s+", " ", text or "")
    return {
        "forestArea": _find_area(t, r"forest area(?:[^0-9]{0,25}(?:proposed|diversion))?")
                       or _find_area(t, r"area proposed for diversion")
                       or _find_area(t, r"\u0935\u0928 \u0915\u094d\u0937\u0947\u0924\u094d\u0930")  # van kshetra
                       or _find_area(t, r"diversion"),
        "totalArea": _find_area(t, r"total(?: project)? area")
                      or _find_area(t, r"\u0915\u0941\u0932 \u092a\u0930\u093f\u092f\u094b\u091c\u0928\u093e"),  # kul pariyojana
        "nonForestArea": _find_area(t, r"non[- ]forest"),
        "caArea": _find_area(t, r"compensatory afforestation") or _find_area(t, r"\bca area"),
        "treeCount": _find_count(t, r"(?:number of |total )?trees?(?:[^0-9]{0,25}(?:felled|felling|enumerat))?")
                      or _find_count(t, r"(?:\u092a\u0947\u095c|\u0935\u0943\u0915\u094d\u0937)"),  # ped / vriksh
    }

def extract_from_table(rows):
    hi = -1
    for i in range(min(len(rows), 6)):
        if rows[i] and any(re.search(r"[a-z\u0900-\u097f]", str(c or "")) for c in rows[i]):
            hi = i
            break
    if hi < 0:
        return {}
    hdr = [str(c or "").lower() for c in rows[hi]]
    def col_of(kws):
        for j, h in enumerate(hdr):
            if any(k in h for k in kws):
                return j
        return -1
    cF = col_of(["forest area", "area (ha)", "area(ha)", "area proposed", "diversion area"])
    cC = col_of(["ca area", "compensatory", "c.a. area"])
    cT = col_of(["tree", "trees"])
    data = rows[hi + 1:]
    def val(c):
        if c < 0:
            return None
        for r in data:
            if r and any(re.search(r"total", str(x or ""), re.I) for x in r):
                v = _num(r[c]) if c < len(r) else None
                if v is not None:
                    return v
        s, any_ = 0.0, False
        for r in data:
            v = _num(r[c]) if r and c < len(r) else None
            if v is not None:
                s += v; any_ = True
        return s if any_ else None
    out = {}
    for k, c in (("forestArea", cF), ("caArea", cC), ("treeCount", cT)):
        v = val(c)
        if v is not None:
            out[k] = v
    return out

# ---------- KML / KMZ geometry ----------
def _ring_area_ha(ring):
    R = 6378137.0
    n = len(ring)
    if n < 3:
        return 0.0
    a = 0.0
    for i in range(n):
        lo1, la1 = ring[i]
        lo2, la2 = ring[(i + 1) % n]
        a += (math.radians(lo2) - math.radians(lo1)) * (2 + math.sin(math.radians(la1)) + math.sin(math.radians(la2)))
    return abs(a * R * R / 2.0) / 10000.0

def _kml_rings(xml):
    rings = []
    for block in re.findall(r"<coordinates>(.*?)</coordinates>", xml, re.S | re.I):
        pts = []
        for tok in block.split():
            parts = tok.split(",")
            if len(parts) >= 2:
                try:
                    pts.append((float(parts[0]), float(parts[1])))
                except ValueError:
                    pass
        if len(pts) >= 3:
            rings.append(pts)
    return rings

def parse_kml_area(xml):
    rings = _kml_rings(xml)
    return sum(_ring_area_ha(r) for r in rings), len(rings)

# ---------- document kind guess ----------
def guess_kind(name, text=""):
    s = (name + " " + (text or "")).lower()
    if re.search(r"site inspection|\bsir\b", s): return "SIR"
    if re.search(r"\bdfo\b|divisional forest officer", s): return "DFO"
    if re.search(r"\bcf\b|conservator of forest", s): return "CF"
    if re.search(r"nodal", s): return "Nodal"
    if re.search(r"enumerat|tree list", s): return "Enumeration"
    if re.search(r"land schedule|khasra|khatauni|schedule", s): return "Schedule"
    if re.search(r"application|part.?i\b|form.?a\b", s): return "Application"
    if name.lower().endswith((".kml", ".kmz")): return "Spatial"
    return ""

# ---------- per-file processing ----------
def process_file(name, data):
    fitz, pytesseract, Image, openpyxl, docx = _lazy()
    ext = (name.rsplit(".", 1)[-1] if "." in name else "").lower()
    rec = {"name": name, "ext": ext, "type": ext.upper() or "?", "mr": False,
           "kind": "", "note": "", "figures": None, "areaHa": None, "scanned": False, "ocr": False}
    try:
        if ext == "kml":
            area, rings = parse_kml_area(data.decode("utf-8", "ignore"))
            rec.update(type="KML", mr=True, areaHa=area, kind="Spatial", note=f"{area:.2f} ha ({rings} polygon)")
        elif ext == "kmz":
            zf = zipfile.ZipFile(io.BytesIO(data))
            kmlname = next((n for n in zf.namelist() if n.lower().endswith(".kml")), None)
            if kmlname:
                area, rings = parse_kml_area(zf.read(kmlname).decode("utf-8", "ignore"))
                rec.update(type="KMZ", mr=True, areaHa=area, kind="Spatial", note=f"{area:.2f} ha")
            else:
                rec.update(type="KMZ", note="No KML inside")
        elif ext in ("xlsx", "xlsm", "xls"):
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            figs, sheets = {}, []
            for ws in wb.worksheets:
                sheets.append(ws.title)
                rows = [list(r) for r in ws.iter_rows(values_only=True)]
                f = extract_from_table(rows)
                for k, v in f.items():
                    figs.setdefault(k, v)
            rec.update(type="Spreadsheet", mr=True, figures=figs, kind=guess_kind(name, " ".join(sheets)),
                       note="Sheets: " + ", ".join(sheets))
        elif ext == "csv":
            import csv as _csv
            rows = list(_csv.reader(io.StringIO(data.decode("utf-8", "ignore"))))
            rec.update(type="CSV", mr=True, figures=extract_from_table(rows), kind=guess_kind(name))
        elif ext == "docx":
            d = docx.Document(io.BytesIO(data))
            text = "\n".join(p.text for p in d.paragraphs)
            for tbl in d.tables:
                for row in tbl.rows:
                    text += "\n" + "\t".join(c.text for c in row.cells)
            rec.update(type="DOCX", mr=bool(text.strip()), figures=extract_figures(text), kind=guess_kind(name, text))
        elif ext == "pdf":
            doc = fitz.open(stream=data, filetype="pdf")
            raw = "".join(p.get_text() for p in doc)
            if len(raw.strip()) >= TEXT_LAYER_MIN:
                rec.update(type="PDF (text)", mr=True, figures=extract_figures(raw),
                           kind=guess_kind(name, raw), note=f"{doc.page_count} page(s)")
            else:
                # SCANNED — rasterize + OCR (English + Hindi)
                ocr = []
                for p in doc:
                    pix = p.get_pixmap(dpi=OCR_DPI)
                    im = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                    ocr.append(pytesseract.image_to_string(im, lang=OCR_LANG))
                text = "\n".join(ocr)
                rec.update(type="PDF (scanned, OCR)", mr=True, ocr=True, scanned=True,
                           figures=extract_figures(text), kind=guess_kind(name, text),
                           note=f"OCR of {doc.page_count} page(s)")
            doc.close()
        elif ext in ("png", "jpg", "jpeg", "webp", "tif", "tiff", "bmp", "gif"):
            im = Image.open(io.BytesIO(data)).convert("RGB")
            text = pytesseract.image_to_string(im, lang=OCR_LANG)
            rec.update(type="Image (OCR)", mr=True, ocr=True, scanned=True,
                       figures=extract_figures(text), kind=guess_kind(name, text), note="OCR of image")
        elif ext == "zip":
            out = []
            zf = zipfile.ZipFile(io.BytesIO(data))
            for n in zf.namelist():
                if n.endswith("/"):
                    continue
                out.extend(process_file(n.rsplit("/", 1)[-1], zf.read(n)))
            return out
        else:
            rec["note"] = "Not auto-readable"
    except Exception as e:
        rec["note"] = f"Could not read: {e}"
    return [rec]

# ---------- GIS intersection (optional; needs shapely) ----------
LAYER_EXTS = (".geojson", ".json", ".kml", ".kmz")

def infer_category(name):
    s = name.lower()
    if "esz" in s or "eco-sensitive" in s or "eco sensitive" in s or "sensitive" in s: return "esz"
    if "tiger" in s or "cth" in s or re.search(r"\btr\b", s) or "reserve" in s: return "tr"
    if "corridor" in s or "\bcor\b" in s: return "cor"
    if "protected" in s or "sanctuary" in s or "national park" in s or "\bpa\b" in s or "\bwls\b" in s or "\bnp\b" in s: return "pa"
    return "other"

def _load_manifest(dirp):
    p = os.path.join(dirp, "manifest.json")
    if os.path.isfile(p):
        try:
            return json.load(open(p, encoding="utf-8"))
        except Exception:
            return {}
    return {}

def list_layers(dirp):
    """Return [{file, category}] for authoritative layers in the fixed folder."""
    if not dirp or not os.path.isdir(dirp):
        return []
    man = _load_manifest(dirp)
    out = []
    for f in sorted(os.listdir(dirp)):
        if f.lower().endswith(LAYER_EXTS):
            cat = man.get(f) or infer_category(f)
            out.append({"file": f, "category": cat})
    return out

def _layer_bytes(dirp, f):
    data = open(os.path.join(dirp, f), "rb").read()
    if f.lower().endswith(".kmz"):
        zf = zipfile.ZipFile(io.BytesIO(data))
        k = next((n for n in zf.namelist() if n.lower().endswith(".kml")), None)
        return zf.read(k) if k else b""
    return data

def gis_check(kml_bytes, layers):
    """layers: list of (name, bytes) GeoJSON or KML. Returns area + intersections."""
    from shapely.geometry import Polygon, shape
    from shapely.ops import unary_union
    rings = _kml_rings(kml_bytes.decode("utf-8", "ignore"))
    if not rings:
        return {"error": "No polygon found in KML"}
    proposal = unary_union([Polygon(r) for r in rings if len(r) >= 3])
    area_ha, _ = parse_kml_area(kml_bytes.decode("utf-8", "ignore"))
    results = []
    for lname, ldata in layers:
        entry = {"layer": lname, "category": infer_category(lname)}
        try:
            geoms = []
            txt = ldata.decode("utf-8", "ignore")
            if lname.lower().endswith((".geojson", ".json")):
                gj = json.loads(txt)
                feats = gj.get("features", [gj]) if isinstance(gj, dict) else []
                for f in feats:
                    g = f.get("geometry") if isinstance(f, dict) and "geometry" in f else f
                    if g:
                        geoms.append(shape(g))
            else:  # KML / KMZ-extracted
                for r in _kml_rings(txt):
                    if len(r) >= 3:
                        geoms.append(Polygon(r))
            layer_geom = unary_union(geoms) if geoms else None
            if layer_geom is None:
                entry["status"] = "no geometry"
            else:
                inter = proposal.intersection(layer_geom)
                hit = (not inter.is_empty) and inter.area > 0
                entry["intersects"] = hit
                entry["overlap_frac"] = round((inter.area / proposal.area) * 100, 2) if hit and proposal.area else 0
        except Exception as e:
            entry["error"] = str(e)
        results.append(entry)
    return {"proposal_area_ha": round(area_ha, 4), "layers": results}

def gis_auto(kml_bytes, dirp):
    """Run intersection against every layer in the fixed folder."""
    meta = list_layers(dirp)
    layers = [(m["file"], _layer_bytes(dirp, m["file"])) for m in meta]
    res = gis_check(kml_bytes, layers)
    cats = {m["file"]: m["category"] for m in meta}
    for r in res.get("layers", []):
        r["category"] = cats.get(r["layer"], r.get("category", "other"))
    res["layers_dir_count"] = len(meta)
    return res
